import wx
import numpy as np
import re
from libraries.Utilities import load_rsf_data
from libraries.Save import save_state
from libraries.Open import load_library_data
from libraries.Sheet_Operations import on_sheet_selected
from libraries.Peak_Functions import AtomicConcentrations


def export_results(window):
    """
    Export peak fitting results to the results grid and update window.Data.
    """
    library_data = load_library_data()
    current_instrument = window.current_instrument

    current_rows = window.results_grid.GetNumberRows()
    start_row = current_rows  # Preserve existing data

    _ensure_results_grid_columns(window)

    peak_data = []
    sheet_name = window.sheet_combobox.GetValue()
    num_peaks = window.peak_params_grid.GetNumberRows() // 2

    for i in range(num_peaks):
        row = i * 2

        peak_params = _extract_peak_parameters(window, row, library_data, current_instrument)
        fitting_model = window.peak_params_grid.GetCellValue(row, 13)

        area, normalized_area, rel_area = _calculate_peak_areas(window, peak_params, row)

        peak_data.append((peak_params['name'], peak_params['position'], peak_params['height'],
                          peak_params['fwhm'], peak_params['lg_ratio'], area, peak_params['rsf'], normalized_area))

        peak_label = _update_data_structure(window, sheet_name, i, peak_params, area, rel_area, fitting_model)

        if start_row + i >= current_rows:
            window.results_grid.AppendRows(1)

        _update_results_grid(window, start_row + i, peak_params, area, rel_area, fitting_model, peak_label)

    window.results_grid.ForceRefresh()
    window.update_checkboxes_from_data()

    _bind_grid_events(window)

    window.update_atomic_percentages()

    save_state(window)

def _ensure_results_grid_columns(window):
    """Ensure that all necessary columns exist in the results grid."""
    if window.results_grid.GetNumberCols() < 11:
        window.results_grid.AppendCols(2)
        window.results_grid.SetColLabelValue(9, "Fitting Model")
        window.results_grid.SetColLabelValue(10, "Rel. Area")


def safe_float(value, default=0.0):
    try:
        return float(value) if value.strip() else default
    except (ValueError, AttributeError):
        return default


def _extract_peak_parameters_OLD(window, row, library_data, current_instrument):
    peak_name = window.peak_params_grid.GetCellValue(row, 1)  # Label

    # Use regex to extract element and orbital information
    match = re.match(r'([A-Z][a-z]*)(\d+[spdf])(?:(\d+/\d+))?', peak_name)
    if match:
        element, orbital, suborbital = match.groups()
        core_level = element + orbital
    else:
        core_level = ''.join(filter(str.isalnum, peak_name.split()[0]))
        element, orbital, suborbital = core_level, '', None

    # Get RSF from library_data
    rsf = get_rsf_from_library(library_data, element, orbital, current_instrument)

    # Adjust RSF for doublets
    if suborbital:
        if orbital.endswith('p'):
            total_electrons = 6
            if suborbital == '3/2':
                rsf *= 4 / total_electrons
            elif suborbital == '1/2':
                rsf *= 2 / total_electrons
        elif orbital.endswith('d'):
            total_electrons = 10
            if suborbital == '5/2':
                rsf *= 6 / total_electrons
            elif suborbital == '3/2':
                rsf *= 4 / total_electrons
        elif orbital.endswith('f'):
            total_electrons = 14
            if suborbital == '7/2':
                rsf *= 8 / total_electrons
            elif suborbital == '5/2':
                rsf *= 6 / total_electrons

    print(f"Element: {element}")
    print(f"Orbital: {orbital}")
    print(f"Base RSF from library: {rsf}")
    print(f"Current instrument: {current_instrument}")

    if suborbital:
        if orbital.endswith('p'):
            total_electrons = 6
            print(f"p orbital - Total electrons: {total_electrons}")
            if suborbital == '3/2':
                print(f"3/2 adjustment: {rsf} * (4/{total_electrons}) = {rsf * 4 / total_electrons}")
            elif suborbital == '1/2':
                print(f"1/2 adjustment: {rsf} * (2/{total_electrons}) = {rsf * 2 / total_electrons}")
        elif orbital.endswith('d'):
            total_electrons = 10
            print(f"d orbital - Total electrons: {total_electrons}")
            if suborbital == '5/2':
                print(f"5/2 adjustment: {rsf} * (6/{total_electrons}) = {rsf * 6 / total_electrons}")
            elif suborbital == '3/2':
                print(f"3/2 adjustment: {rsf} * (4/{total_electrons}) = {rsf * 4 / total_electrons}")

    return {
        'name': peak_name,
        'position': safe_float(window.peak_params_grid.GetCellValue(row, 2)),
        'height': safe_float(window.peak_params_grid.GetCellValue(row, 3)),
        'fwhm': safe_float(window.peak_params_grid.GetCellValue(row, 4)),
        'lg_ratio': safe_float(window.peak_params_grid.GetCellValue(row, 5)),
        'rsf': rsf,
        'area': safe_float(window.peak_params_grid.GetCellValue(row, 6)),
        'sigma': safe_float(window.peak_params_grid.GetCellValue(row, 7)),
        'gamma': safe_float(window.peak_params_grid.GetCellValue(row, 8)),
        'skew': safe_float(window.peak_params_grid.GetCellValue(row, 9)),
        'constraints': {
            'position': window.peak_params_grid.GetCellValue(row + 1, 2),
            'height': window.peak_params_grid.GetCellValue(row + 1, 3),
            'fwhm': window.peak_params_grid.GetCellValue(row + 1, 4),
            'lg_ratio': window.peak_params_grid.GetCellValue(row + 1, 5),
            'area': window.peak_params_grid.GetCellValue(row+1, 6),
            'sigma': window.peak_params_grid.GetCellValue(row+1, 7),
            'gamma': window.peak_params_grid.GetCellValue(row+1, 8),
            'skew': window.peak_params_grid.GetCellValue(row+1, 9)
        }
    }


def _extract_peak_parameters(window, row, library_data, current_instrument):
    peak_name = window.peak_params_grid.GetCellValue(row, 1)  # Label

    # Use regex to extract element, orbital, and suborbital
    match = re.match(r'([A-Z][a-z]*)(\d+[spdf])(?:(\d+/\d+))?', peak_name)
    if match:
        element, orbital, suborbital = match.groups()
        if suborbital:
            # Use full orbital including suborbital for RSF lookup
            core_level = f"{element}{orbital}{suborbital}"
        else:
            core_level = f"{element}{orbital}"
    else:
        core_level = ''.join(filter(str.isalnum, peak_name.split()[0]))
        element, orbital, suborbital = core_level, '', None

    # Get RSF directly using complete orbital designation
    key = (element, orbital + (suborbital or ''))
    if key in library_data and current_instrument in library_data[key]:
        rsf = library_data[key][current_instrument]['rsf']
    else:
        # Fallback to Al if instrument not found
        rsf = library_data[key]['Al']['rsf'] if key in library_data else 1.0

    return {
        'name': peak_name,
        'position': safe_float(window.peak_params_grid.GetCellValue(row, 2)),
        'height': safe_float(window.peak_params_grid.GetCellValue(row, 3)),
        'fwhm': safe_float(window.peak_params_grid.GetCellValue(row, 4)),
        'lg_ratio': safe_float(window.peak_params_grid.GetCellValue(row, 5)),
        'rsf': rsf,
        'area': safe_float(window.peak_params_grid.GetCellValue(row, 6)),
        'sigma': safe_float(window.peak_params_grid.GetCellValue(row, 7)),
        'gamma': safe_float(window.peak_params_grid.GetCellValue(row, 8)),
        'skew': safe_float(window.peak_params_grid.GetCellValue(row, 9)),
        'constraints': {
            'position': window.peak_params_grid.GetCellValue(row + 1, 2),
            'height': window.peak_params_grid.GetCellValue(row + 1, 3),
            'fwhm': window.peak_params_grid.GetCellValue(row + 1, 4),
            'lg_ratio': window.peak_params_grid.GetCellValue(row + 1, 5),
            'area': window.peak_params_grid.GetCellValue(row+1, 6),
            'sigma': window.peak_params_grid.GetCellValue(row+1, 7),
            'gamma': window.peak_params_grid.GetCellValue(row+1, 8),
            'skew': window.peak_params_grid.GetCellValue(row+1, 9)
        }
    }

def get_rsf_from_library_OLD(library_data, element, orbital, instrument):
    key = (element, orbital)
    if key in library_data and instrument in library_data[key]:
        return library_data[key][instrument]['rsf']

    return library_data[key]['Al1486']['rsf']  # Default to Al if not found


def get_rsf_from_library(library_data, element, orbital, instrument):
   key = (element, orbital)
   if key in library_data and instrument in library_data[key]:
       print(f"Found {element}{orbital} at row {library_data[key][instrument]['row']}")
       return library_data[key][instrument]['rsf']
   return library_data[key]['Al']['rsf']

def _calculate_peak_areas(window, peak_params, row):
    area = float(window.peak_params_grid.GetCellValue(row, 6))
    peak_name = window.peak_params_grid.GetCellValue(row, 0)
    rsf = peak_params['rsf']

    binding_energy = float(window.peak_params_grid.GetCellValue(row, 2))
    kinetic_energy = window.photons - binding_energy

    if window.library_type == "Scofield":
        ecf = kinetic_energy ** 0.6
    elif window.library_type == "Wagner":
        ecf = kinetic_energy ** 1.0
    elif window.library_type == "TPP-2M":
        # Calculate IMFP using TPP-2M using the average matrix
        imfp = AtomicConcentrations.calculate_imfp_tpp2m(kinetic_energy)

        # 26.2 is a factor added by Avantage to match KE^0.6
        ecf = imfp * 26.2
    elif window.library_type == "EAL":
        z_avg = 50  # Default values

        eal = (0.65 + 0.007 * kinetic_energy**0.93) / (z_avg**0.38)
        ecf = eal
    elif window.library_type == "None":
        ecf = 1.0
    else:
        ecf = 1.0

    # Angular correction
    angular_correction = 1.0
    if window.use_angular_correction:
        orbital_type = peak_params['name'][-1].lower()  # Get orbital type from name
        angular_correction = AtomicConcentrations.calculate_angular_correction(
            window,
            peak_name,
            window.analysis_angle
        )

    txfn = 1.0
    normalized_area = area / (rsf * txfn * ecf * angular_correction)
    rel_area = normalized_area

    return round(area, 2), round(normalized_area, 2), round(rel_area, 2)

def _update_results_grid(window, row, peak_params, area, rel_area, fitting_model, peak_label):
    """Update a row in the results grid with peak data."""
    window.results_grid.SetCellValue(row, 0, f"{peak_params['name']}")  # Keep the original peak name
    window.results_grid.SetCellValue(row, 1, f"{peak_params['position']:.2f}")
    window.results_grid.SetCellValue(row, 2, f"{peak_params['height']:.2f}")
    window.results_grid.SetCellValue(row, 3, f"{peak_params['fwhm']:.2f}")
    window.results_grid.SetCellValue(row, 4, f"{peak_params['lg_ratio']:.2f}")
    window.results_grid.SetCellValue(row, 5, f"{area:.2f}")
    window.results_grid.SetCellValue(row, 6, "0.00")  # Initial atomic percentage

    checkbox_state = window.Data['Results']['Peak'][peak_label].get('Checkbox', '0')
    _set_checkbox(window, row, 7, checkbox_state)

    window.results_grid.SetCellValue(row, 8, f"{peak_params['rsf']:.2f}")
    window.results_grid.SetCellValue(row, 9, "1.0")  # TXFN default value
    print(f"Library: {window.library_type}")
    if window.library_type == "Scofield":
        window.results_grid.SetCellValue(row, 10, "KE^0.6")
    elif window.library_type == "Wagner":
        window.results_grid.SetCellValue(row, 10, "KE^1.0")
    elif window.library_type == "TPP-2M":
        window.results_grid.SetCellValue(row, 10, "TPP-2M")
    elif window.library_type == "EAL":
        window.results_grid.SetCellValue(row, 10, "EAL")
    else:
        window.results_grid.SetCellValue(row, 10, "1.0")


    window.results_grid.SetCellValue(row, 11, window.current_instrument)
    window.results_grid.SetCellValue(row, 12, fitting_model)
    window.results_grid.SetCellValue(row, 13, f"{rel_area:.2f}")
    window.results_grid.SetCellValue(row, 14, f"{peak_params['sigma']:.2f}")  # Sigma
    window.results_grid.SetCellValue(row, 15, f"{peak_params['gamma']:.2f}")  # Gamma
    window.results_grid.SetCellValue(row, 17, f"{window.bg_min_energy:.2f}" if window.bg_min_energy is not None else "")
    window.results_grid.SetCellValue(row, 18, f"{window.bg_max_energy:.2f}" if window.bg_max_energy is not None else "")
    window.results_grid.SetCellValue(row, 21, window.sheet_combobox.GetValue())
    _set_constraints(window, row, peak_params['constraints'])

    # Force a refresh of the grid cell to ensure the checkbox is displayed correctly
    window.results_grid.RefreshAttr(row, 7)


def _set_checkbox(window, row, col, state='0'):
    """Set up a checkbox in the specified grid cell."""
    window.results_grid.SetCellRenderer(row, col, wx.grid.GridCellBoolRenderer())
    window.results_grid.SetCellEditor(row, col, wx.grid.GridCellBoolEditor())
    window.results_grid.SetCellValue(row, col, state)
    # window.results_grid.SetCellValue(row, col, '1' if state == '1' else '0')
    window.results_grid.ForceRefresh()

def _set_constraints(window, row, constraints):
    """Set constraint values in the results grid."""
    window.results_grid.SetCellValue(row, 22, constraints['position'])
    window.results_grid.SetCellValue(row, 23, constraints['height'])
    window.results_grid.SetCellValue(row, 24, constraints['fwhm'])
    window.results_grid.SetCellValue(row, 25, constraints['lg_ratio'])
    window.results_grid.SetCellValue(row, 26, constraints['area'])
    window.results_grid.SetCellValue(row, 27, constraints['sigma'])
    window.results_grid.SetCellValue(row, 28, constraints['gamma'])


def _update_data_structure(window, sheet_name, peak_index, peak_params, area, rel_area, fitting_model):
    """Update the window.Data structure with peak results."""
    results = window.Data['Results']['Peak']

    # Find the next available peak number
    existing_peaks = [int(key.split('_')[1]) for key in results.keys() if key.startswith('Peak_')]
    next_peak_number = max(existing_peaks + [-1]) + 1

    peak_label = f"Peak_{next_peak_number}"
    peak_name = peak_params['name']

    # Check if this peak already exists (by name and sheet)
    existing_peak = next((key for key, value in results.items()
                          if value['Name'] == peak_name and value['Sheetname'] == sheet_name), None)

    if existing_peak:
        peak_label = existing_peak

    if window.library_type == "Scofield":
        ecf_type = "KE^0.6"
    elif window.library_type == "Wagner":
        ecf_type = "KE^1.0"
    elif window.library_type == "TPP-2M":
        ecf_type = "TPP-2M"
    else:
        ecf_type = "1.0"

    peak_data = {
        'Label': peak_label,
        'Name': peak_name,
        'Position': peak_params['position'],
        'Height': peak_params['height'],
        'FWHM': peak_params['fwhm'],
        'L/G': peak_params['lg_ratio'],
        'Area': area,
        'at. %': results.get(peak_label, {}).get('at. %', 0.00),  # Preserve existing at. % if available
        'RSF': peak_params['rsf'],
        'TXFN': 1.0,
        'ECF': ecf_type,
        'Instrument':  window.current_instrument,
        'Fitting Model': fitting_model,
        'Rel. Area': rel_area,
        'Sigma': peak_params['sigma'],
        'Gamma': peak_params['gamma'],
        'Skew': peak_params.get('skew', 0.0),  # Default to 0.0 if 'skew' is not in peak_params
        'Bkg Low': window.bg_min_energy,
        'Bkg High': window.bg_max_energy,
        'Sheetname': sheet_name,
        'Pos. Constraint': peak_params['constraints']['position'],
        'Height Constraint': peak_params['constraints']['height'],
        'FWHM Constraint': peak_params['constraints']['fwhm'],
        'L/G Constraint': peak_params['constraints']['lg_ratio'],
        'Area Constraint': peak_params['constraints']['area'],
        'Sigma Constraint':peak_params['constraints']['sigma'],
        'Gamma Constraint':peak_params['constraints']['gamma'],
        'Checkbox': results.get(peak_label, {}).get('Checkbox', '0')  # Preserve existing checkbox state if available
    }

    window.Data['Results']['Peak'][peak_label] = peak_data
    return peak_label  # Return the peak label for reference


def _bind_grid_events(window):
    """Bind necessary events to the results grid."""
    window.results_grid.Bind(wx.grid.EVT_GRID_CELL_CHANGED, window.on_cell_changed)


def export_word_report(window):
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.section import WD_ORIENT
    import os
    from docx.shared import RGBColor

    file_path = window.Data['FilePath']
    base_path = os.path.splitext(file_path)[0]
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    doc.add_heading('XPS Analysis Report', 0)

    for sheet_name in window.Data['Core levels'].keys():
        window.sheet_combobox.SetValue(sheet_name)
        on_sheet_selected(window, sheet_name)

        doc.add_heading(sheet_name, level=1)

        plot_path = f"{base_path}_{sheet_name}_plot.png"

        is_survey = "survey" in sheet_name.lower() or "wide" in sheet_name.lower()
        width = window.survey_word_width if is_survey else window.word_width
        height = window.survey_word_height if is_survey else window.word_height
        dpi = window.survey_word_dpi if is_survey else window.word_dpi

        original_size = window.figure.get_size_inches()
        window.figure.set_size_inches(width, height)
        window.figure.savefig(plot_path, dpi=dpi, bbox_inches='tight')
        window.figure.set_size_inches(original_size)

        doc.add_picture(plot_path, width=Inches(7))
        os.remove(plot_path)

        doc.add_heading('Fitting Parameters', level=2)

        # Peak parameters table
        columns = ['Label', 'Position', 'FWHM', 'Area', 'L/G', 'Sigma', 'Gamma', 'Model']
        col_indices = [1, 2, 4, 6, 5, 7, 8, 13]
        col_widths = [1.5, 1., 1., 1.5, 1., 1., 1., 2]

        num_peaks = window.peak_params_grid.GetNumberRows() // 2

        table = doc.add_table(rows=1 + num_peaks * 2, cols=len(columns))
        table.style = 'Table Grid'
        table.autofit = False

        # Add headers with bold formatting
        for i, header in enumerate(columns):
            cell = table.cell(0, i)
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        for peak in range(num_peaks):
            value_row = peak * 2 + 1
            constraint_row = value_row + 1
            grid_row = peak * 2

            for col, grid_col in enumerate(col_indices):
                cell = table.cell(value_row, col)
                cell.text = window.peak_params_grid.GetCellValue(grid_row, grid_col)

            for col, grid_col in enumerate(col_indices):
                if col > 0:
                    cell = table.cell(constraint_row, col)
                    cell.text = window.peak_params_grid.GetCellValue(grid_row + 1, grid_col)

        for i, col in enumerate(table.columns):
            col.width = Inches(col_widths[i])

        doc.add_paragraph('')

    doc.add_paragraph('')

    # Results grid table
    doc.add_heading('Quantification Results', level=2)

    results_columns = ['Peak Label', 'Position', 'FWHM', 'Area', 'Rel. Area', 'RSF', 'at. %']
    results_indices = [0, 1, 3, 5, 10, 8, 6]  # Indices in results grid
    results_widths = [2, 1., 1., 1.5, 1.5, 0.9, 1.]

    results_table = doc.add_table(rows=window.results_grid.GetNumberRows() + 1, cols=len(results_columns))
    results_table.style = 'Table Grid'
    results_table.autofit = False

    # Add headers with bold formatting
    for i, header in enumerate(results_columns):
        cell = results_table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Add data
    for row in range(window.results_grid.GetNumberRows()):
        for col, grid_col in enumerate(results_indices):
            cell = results_table.cell(row + 1, col)
            cell.text = window.results_grid.GetCellValue(row, grid_col)

    # Set column widths
    for i, col in enumerate(results_table.columns):
        col.width = Inches(results_widths[i])



    word_path = f"{base_path}_report.docx"
    doc.save(word_path)
    window.show_popup_message2("Report Created", f"Saved to: {word_path}")